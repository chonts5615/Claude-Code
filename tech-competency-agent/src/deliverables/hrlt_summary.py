"""HRLT Summary writer — 1-page docx tuned for a 90-second executive read.

No IO/HR jargon. Sections, in order:

1. Title (Georgia, 18pt, Cargill Leaf Green)
2. Executive bullet (3 sentences max)
3. Top 3 findings
4. Coverage statistic
5. Next-step recommendation
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

from src.utils.branding import ARIAL, BLACK, GEORGIA, LEAF_GREEN


def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _add_title(doc: Document, family: str, run_id: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{family} — Technical Competency Snapshot")
    run.font.name = GEORGIA
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = _hex_to_rgb(LEAF_GREEN)

    sub = doc.add_paragraph()
    sub_run = sub.add_run(f"Run ID: {run_id}")
    sub_run.font.name = ARIAL
    sub_run.font.size = Pt(9)
    sub_run.font.italic = True
    sub_run.font.color.rgb = _hex_to_rgb(BLACK)


def _add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = GEORGIA
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = _hex_to_rgb(LEAF_GREEN)


def _add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = ARIAL
    run.font.size = Pt(10)
    run.font.color.rgb = _hex_to_rgb(BLACK)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = ARIAL
    run.font.size = Pt(10)
    run.font.color.rgb = _hex_to_rgb(BLACK)


def _executive_bullet(family: str, total: int, coverage_pct: float) -> str:
    return (
        f"The {family} family now has {total} technical competencies in the master library. "
        f"Job coverage is at {coverage_pct:.0f}%. "
        "The library is ready for SME validation pending the noted next step."
    )


def _next_step(coverage_pct: float) -> str:
    if coverage_pct >= 90.0:
        return "Proceed to SME R2 review and lock the family for FINAL stage."
    return (
        "Close the coverage gap before SME review — re-run extraction on the "
        "uncovered essential functions and re-validate against the BCO ledger."
    )


def write_hrlt_summary(
    out_path: Path,
    family: str,
    top_findings: list[str],
    coverage_pct: float,
    total_competencies: int,
    run_id: str,
) -> Path:
    """Write the 1-page HRLT summary docx.

    Args:
        out_path: Output docx path (parents created if missing).
        family: Job family.
        top_findings: 3 short finding statements (extras are truncated, gaps padded).
        coverage_pct: Job coverage percentage (0-100).
        total_competencies: Library competency count for the family.
        run_id: Run identifier.

    Returns:
        The resolved output path.
    """
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    _add_title(doc, family=family, run_id=run_id)

    _add_section_heading(doc, "Executive summary")
    _add_body(doc, _executive_bullet(family, total_competencies, coverage_pct))

    _add_section_heading(doc, "Top findings")
    findings = list(top_findings[:3])
    while len(findings) < 3:
        findings.append("(no additional finding)")
    for f in findings:
        _add_bullet(doc, f)

    _add_section_heading(doc, "Coverage")
    _add_body(
        doc,
        f"{coverage_pct:.1f}% of essential functions in {family} are covered "
        f"by the {total_competencies} technical competencies in the library.",
    )

    _add_section_heading(doc, "Recommended next step")
    _add_body(doc, _next_step(coverage_pct))

    doc.save(out_path)
    return out_path
