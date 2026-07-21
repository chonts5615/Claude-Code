#!/usr/bin/env python3
"""Round 2 SME Review Packet generator.

Renders a fully formatted, Cargill-branded DOCX packet from a single JSON
data file following the locked Round 2 SME Focus Group Packet format
(see CLAUDE.md). Structure, taxonomy, and visual system are fixed; only the
content (names, competencies, quotes, counts, coverage cells, dispositions)
changes per sub-family.

Usage:
    python generator.py <data.json> [output.docx]
    python generator.py --all          # render every data/*.json to outputs/
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# --- Cargill brand colors (locked) -----------------------------------------
LEAF_GREEN = RGBColor(0, 132, 61)  # #00843D  confirmation / brand
DEEP_GREEN = RGBColor(1, 99, 45)  # #01632D
PROPOSED_BLUE = RGBColor(15, 73, 197)  # #0F49C5  proposed
AMBER = RGBColor(254, 168, 0)  # #FEA800  needs decision
ALERT_RED = RGBColor(197, 15, 31)  # #C50F1F  needs quick answer
GRAY = RGBColor(112, 119, 115)  # #707773  kept / muted ink
INK = RGBColor(28, 39, 34)  # #1C2722  body
WHITE = RGBColor(255, 255, 255)

# hex fills (python-docx shading takes hex strings)
TINT = "F5F9ED"  # white green
TINT_SOFT = "FCFDFB"
HEAD_GREEN = "00843D"
ROW_HEAD = "EDEDED"

# Status taxonomy (locked) -> (badge color, fill hex)
STATUS_STYLE = {
    "PROPOSED": (PROPOSED_BLUE, "E7EDFB"),
    "NEEDS YOUR DECISION": (AMBER, "FFF3DA"),
    "NEEDS YOUR CONFIRMATION": (LEAF_GREEN, "E3F1E8"),
    "KEPT, NO ACTION": (GRAY, "EEF0EF"),
    "NEEDS A QUICK ANSWER": (ALERT_RED, "FBE3E5"),
}

BODY_FONT = "Arial"
SERIF_FONT = "Georgia"  # display / headings


# --- low-level helpers ------------------------------------------------------
def shade(cell, fill_hex: str) -> None:
    """Apply a solid background fill to a table cell."""
    tcpr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcpr.append(shd)


def set_repeat_header(row) -> None:
    """Mark a table row to repeat as a header on each page."""
    tr = row._tr
    trpr = tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trpr.append(th)


def run(p, text, *, bold=False, italic=False, size=11, color=None, font=BODY_FONT):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return r


def para(doc, text="", *, bold=False, italic=False, size=11, color=None, space_after=6,
         space_before=0, font=BODY_FONT, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        p.alignment = align
    if text:
        run(p, text, bold=bold, italic=italic, size=size, color=color, font=font)
    return p


def lead_para(doc, lead, text, *, size=11):
    """A paragraph that opens with a bold lead phrase, then continues."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run(p, lead + "  ", bold=True, size=size, color=INK)
    run(p, text, size=size, color=INK)
    return p


def heading(doc, text, *, size=16, color=DEEP_GREEN, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run(p, text, bold=True, size=size, color=color, font=SERIF_FONT)
    return p


def styled_table(doc, n_rows, n_cols, *, header=True, widths=None):
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if header and n_rows:
        set_repeat_header(table.rows[0])
    return table


def fill_cell(cell, text, *, bold=False, italic=False, size=10, color=INK,
              fill=None, align=None, font=BODY_FONT):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    if align is not None:
        p.alignment = align
    if text:
        run(p, text, bold=bold, italic=italic, size=size, color=color, font=font)
    if fill:
        shade(cell, fill)


def header_row(table, headers, *, fill=HEAD_GREEN):
    for i, h in enumerate(headers):
        fill_cell(table.rows[0].cells[i], h, bold=True, size=10, color=WHITE, fill=fill)


# --- section renderers ------------------------------------------------------
def render_cover(doc, m):
    para(doc, m.get("eyebrow", "Technical Competency"), bold=True, size=10,
         color=LEAF_GREEN, space_after=2)
    heading(doc, m.get("title", "Round 2 SME Review Packet"), size=26, color=DEEP_GREEN,
            space_before=0, space_after=4)
    para(doc, m.get("sub_family", ""), bold=True, size=15, color=INK, space_after=2,
         font=SERIF_FONT)
    para(doc, m.get("subtitle", ""), italic=True, size=11, color=GRAY, space_after=10)
    para(doc, f"Function: {m.get('function', '')}", size=10, color=INK, space_after=1)
    para(doc, f"Prepared by: {m.get('prepared_by', '')}", size=10, color=INK, space_after=1)
    para(doc, f"Date: {m.get('date', '')}", size=10, color=INK, space_after=8)

    # confidentiality band
    t = styled_table(doc, 1, 1, header=False)
    fill_cell(
        t.rows[0].cells[0],
        m.get(
            "confidentiality",
            "Proprietary competency architecture developed by Cargill's Assessment, "
            "Competency & Career Framework team. Limited to designated SME reviewers and "
            "project stakeholders. Do not forward or share outside the approved panel "
            "without written authorization from the project owner.",
        ),
        italic=True, size=9, color=GRAY, fill=TINT,
    )


def render_reviewer_info(doc, m):
    heading(doc, "Reviewer information")
    rows = [
        ("Reviewer name", ""),
        ("Specialization / group", m.get("specialization_label", m.get("sub_family", ""))),
        ("Return by", m.get("return_by", "[RETURN_BY_DATE]")),
        ("Time required", m.get("time_required", "About 35 minutes for this focus group")),
    ]
    t = styled_table(doc, len(rows), 2, header=False)
    for i, (k, v) in enumerate(rows):
        fill_cell(t.rows[i].cells[0], k, bold=True, size=10, color=INK, fill=ROW_HEAD)
        fill_cell(t.rows[i].cells[1], v, size=10, color=INK)


def render_how_it_works(doc, hw, status_counts):
    heading(doc, "How this packet works")
    lead_para(doc, "What we did.", hw.get("what_we_did", ""))
    lead_para(doc, "What we need from you.", hw.get("what_we_need", ""))
    lead_para(doc, "Where your feedback went.", hw.get("where_feedback_went", ""))
    if hw.get("status_intro"):
        para(doc, hw["status_intro"], size=11, color=INK)

    # status legend table with live counts
    headers = ["Status", "What it means", "What to do", "You have"]
    legend = [
        ("PROPOSED", "A new or expanded competency, drafted in full.",
         "Review and react; confirm it fits or note changes."),
        ("NEEDS YOUR DECISION",
         "Reviewers saw this differently, or it changes who owns the work. A recommendation is drafted.",
         "Review the recommendation; bring your view to the focus group."),
        ("NEEDS YOUR CONFIRMATION",
         "We applied your feedback or a clear, agreed change. Structure is unchanged.",
         "Read it; confirm or note a correction."),
        ("KEPT, NO ACTION", "Strongly supported and unchanged.",
         "Acknowledge; nothing to do."),
        ("NEEDS A QUICK ANSWER", "One specific question only you can answer.",
         "Reply in the box provided."),
    ]
    t = styled_table(doc, len(legend) + 1, 4)
    header_row(t, headers)
    for i, (status, meaning, todo) in enumerate(legend, 1):
        badge_color, _ = STATUS_STYLE[status]
        fill_cell(t.rows[i].cells[0], status, bold=True, size=9, color=badge_color)
        fill_cell(t.rows[i].cells[1], meaning, size=9, color=INK)
        fill_cell(t.rows[i].cells[2], todo, size=9, color=INK)
        fill_cell(t.rows[i].cells[3], str(status_counts.get(status, 0)), bold=True,
                  size=10, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)


def render_ground_rules(doc, rules):
    heading(doc, "Ground rules for the session")
    for rule in rules:
        lead_para(doc, rule.get("lead", ""), rule.get("text", ""))


def render_start_here(doc, sh):
    heading(doc, sh.get("heading", "Start here"))
    if sh.get("intro"):
        para(doc, sh["intro"], size=11, color=INK)

    tbl = sh.get("table")
    if tbl:
        headers = tbl["headers"]
        t = styled_table(doc, len(tbl["rows"]) + 1, len(headers))
        header_row(t, headers)
        for i, row in enumerate(tbl["rows"], 1):
            for j, val in enumerate(row):
                fill_cell(t.rows[i].cells[j], val, size=9, color=INK,
                          bold=(j == 0))

    quote = sh.get("quote")
    if quote:
        q = styled_table(doc, 1, 1, header=False)
        cell = q.rows[0].cells[0]
        shade(cell, TINT)
        p = cell.paragraphs[0]
        run(p, quote.get("label", "") + "  ", bold=True, size=9, color=DEEP_GREEN)
        run(p, "“" + quote.get("text", "") + "”", italic=True, size=9, color=INK)

    if sh.get("decide"):
        lead_para(doc, "Decide.", sh["decide"])


def render_coverage(doc, cov):
    heading(doc, "Specialization coverage matrix")
    specs = cov["specializations"]
    headers = ["#", "Technical competency"] + specs
    rows = cov["rows"]
    t = styled_table(doc, len(rows) + 1, len(headers))
    header_row(t, headers)
    for i, r in enumerate(rows, 1):
        fill_cell(t.rows[i].cells[0], r.get("num", str(i)), size=9, color=INK,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(t.rows[i].cells[1], r["competency"], size=9, color=INK, bold=True)
        for j, role in enumerate(r.get("roles", []), 2):
            fill_cell(t.rows[i].cells[j], role, size=9, color=DEEP_GREEN, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Legend:  " + cov.get("legend", ""), italic=True, size=8.5, color=GRAY,
         space_before=4)
    if cov.get("note"):
        para(doc, cov["note"], size=9, color=INK)


def render_ratings_legend(doc, rl):
    heading(doc, rl.get("heading", "How to read the Round 1 ratings"))
    headers = rl.get("headers", ["Score", "Label", "Definition"])
    t = styled_table(doc, len(rl["rows"]) + 1, len(headers))
    header_row(t, headers)
    for i, row in enumerate(rl["rows"], 1):
        for j, val in enumerate(row):
            fill_cell(t.rows[i].cells[j], val, size=9, color=INK, bold=(j <= 1))


def render_status_strip(doc, status, status_note, assigned_to):
    badge_color, fill = STATUS_STYLE.get(status, (GRAY, ROW_HEAD))
    t = styled_table(doc, 1, 2, header=False)
    left = t.rows[0].cells[0]
    shade(left, fill)
    p = left.paragraphs[0]
    run(p, status, bold=True, size=10, color=badge_color)
    if status_note:
        run(p, "  —  " + status_note, size=9, color=INK)
    fill_cell(t.rows[0].cells[1], "Assigned to:  " + assigned_to, size=9, color=INK)


def render_indicators(doc, indicators):
    t = styled_table(doc, len(indicators) + 1, 2)
    header_row(t, ["Level", "Behavioral indicators  (Level 1 basic to Level 4 expert)"])
    for i, ind in enumerate(indicators, 1):
        label = ind["level"] + ("  · updated" if ind.get("updated") else "")
        color = DEEP_GREEN if ind["level"] in ("L3", "L4") else INK
        fill_cell(t.rows[i].cells[0], label, bold=True, size=9, color=color,
                  fill=(TINT if ind["level"] in ("L3", "L4") else None))
        fill_cell(t.rows[i].cells[1], ind["text"], size=9, color=INK)


def render_prompt_box(doc, lead, text):
    t = styled_table(doc, 1, 1, header=False)
    cell = t.rows[0].cells[0]
    shade(cell, TINT_SOFT)
    p = cell.paragraphs[0]
    run(p, lead + "  ", bold=True, size=9, color=DEEP_GREEN)
    run(p, text, size=9, color=INK)


def render_part_a(doc, part_a):
    heading(doc, "Part A · Modifications already applied (read and confirm)")
    if part_a.get("intro"):
        para(doc, part_a["intro"], size=11, color=INK)
    for item in part_a.get("items", []):
        render_status_strip(doc, item["status"], item.get("status_note", ""),
                             item.get("assigned_to", ""))
        heading(doc, item["title"], size=13, color=INK, space_before=6, space_after=4)
        if item.get("what_changed"):
            lead_para(doc, "What changed, and what it means for you.", item["what_changed"])
        if item.get("round1_signal"):
            lead_para(doc, "Round 1 signal.", item["round1_signal"])
        if item.get("indicators"):
            render_indicators(doc, item["indicators"])
        action = item.get("action")
        if action:
            render_prompt_box(doc, action.get("lead", "Confirm."), action.get("text", ""))


def render_what_each_group_said(doc, weg):
    para(doc, "What each group said", bold=True, size=11, color=DEEP_GREEN, space_before=4,
         space_after=2)
    headers = weg["headers"]
    t = styled_table(doc, len(weg["rows"]) + 1, len(headers))
    header_row(t, headers)
    for i, row in enumerate(weg["rows"], 1):
        for j, val in enumerate(row):
            fill_cell(t.rows[i].cells[j], "“" + val + "”", italic=True, size=9,
                      color=INK)


def render_part_b(doc, part_b):
    heading(doc, "Part B · For the focus group (decisions and proposals)")
    if part_b.get("intro"):
        para(doc, part_b["intro"], size=11, color=INK)
    for item in part_b.get("items", []):
        render_status_strip(doc, item["status"], item.get("status_note", ""),
                             item.get("assigned_to", ""))
        title = (item.get("number", "") + "  " + item["title"]).strip()
        heading(doc, title, size=13, color=INK, space_before=6, space_after=2)
        if item.get("owners_line"):
            para(doc, item["owners_line"], italic=True, size=9, color=GRAY, space_after=4)
        note = item.get("note")
        if note:
            render_prompt_box(doc, note.get("lead", "Note."), note.get("text", ""))
        if item.get("what_each_group_said"):
            render_what_each_group_said(doc, item["what_each_group_said"])
        if item.get("recommendation"):
            lead_para(doc, "Recommendation.", item["recommendation"])
        decide = item.get("decide")
        if decide:
            render_prompt_box(doc, decide.get("lead", "Decide."), decide.get("text", ""))


def render_boundaries(doc, b):
    heading(doc, "Boundaries to note")
    headers = b.get("headers", ["Item", "What Round 1 said", "Where it goes"])
    t = styled_table(doc, len(b["rows"]) + 1, len(headers))
    header_row(t, headers)
    for i, row in enumerate(b["rows"], 1):
        for j, val in enumerate(row):
            fill_cell(t.rows[i].cells[j], val, size=9, color=INK, bold=(j == 0))


def render_sign_off(doc, so):
    heading(doc, "Decision capture and sign-off")
    if so.get("intro"):
        para(doc, so["intro"], size=11, color=INK)
    headers = ["Item", "Decision (assignment, level, owner)", "Owner", "Date"]
    rows = so["rows"]
    t = styled_table(doc, len(rows) + 1, 4)
    header_row(t, headers)
    for i, item in enumerate(rows, 1):
        fill_cell(t.rows[i].cells[0], item, size=9, color=INK)
        for j in (1, 2, 3):
            fill_cell(t.rows[i].cells[j], "", size=9, color=INK)


DISPOSITION_COLOR = {
    "Applied": LEAF_GREEN,
    "Synthesized / Used": PROPOSED_BLUE,
    "Via level": GRAY,
    "Deferred": AMBER,
    "Corrected": ALERT_RED,
    "Kept": GRAY,
}


def render_disposition_register(doc, dr):
    doc.add_page_break()
    heading(doc, "Reference appendix")
    para(doc, "Everything below is reference. Appendix A records where every Round 1 "
              "comment went; Appendix B holds the full competency indicators.",
         italic=True, size=9, color=GRAY)
    heading(doc, "Appendix A · Disposition register (where your feedback went)", size=14,
            color=DEEP_GREEN, space_before=10)
    if dr.get("intro"):
        para(doc, dr["intro"], size=10, color=INK)
    headers = ["Competency", "What Round 1 said", "Disposition", "Outcome"]
    entries = dr.get("entries", [])
    t = styled_table(doc, len(entries) + 1, 4)
    header_row(t, headers)
    for i, e in enumerate(entries, 1):
        fill_cell(t.rows[i].cells[0], e.get("competency", ""), bold=True, size=9, color=INK)
        fill_cell(t.rows[i].cells[1], e.get("round1", ""), size=9, color=INK)
        disp = e.get("disposition", "")
        fill_cell(t.rows[i].cells[2], disp, bold=True, size=9,
                  color=DISPOSITION_COLOR.get(disp, INK))
        fill_cell(t.rows[i].cells[3], e.get("outcome", ""), size=9, color=INK)


def render_appendix(doc, ap):
    doc.add_page_break()
    heading(doc, "Appendix B · Competency definitions and indicators for reference", size=14,
            color=DEEP_GREEN)
    if ap.get("intro"):
        para(doc, ap["intro"], size=10, color=GRAY)
    for comp in ap.get("competencies", []):
        heading(doc, comp["name"], size=13, color=DEEP_GREEN, space_before=10, space_after=2)
        if comp.get("note"):
            para(doc, comp["note"], italic=True, size=9, color=GRAY)
        if comp.get("indicators"):
            render_indicators(doc, comp["indicators"])


def render_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, "Proprietary — SME reviewers and project stakeholders only  |  "
           "Assessment, Competency & Career Framework Team · 2026",
        size=8, color=GRAY)


# --- orchestration ----------------------------------------------------------
def compute_status_counts(data):
    counts = {k: 0 for k in STATUS_STYLE}
    for item in data.get("part_a", {}).get("items", []):
        counts[item["status"]] = counts.get(item["status"], 0) + 1
    for item in data.get("part_b", {}).get("items", []):
        counts[item["status"]] = counts.get(item["status"], 0) + 1
    return counts


def build_document(data) -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Pt(54)
        sec.left_margin = sec.right_margin = Pt(54)

    counts = compute_status_counts(data)

    render_cover(doc, data["metadata"])
    render_reviewer_info(doc, data["metadata"])
    render_how_it_works(doc, data.get("how_this_works", {}), counts)
    render_ground_rules(doc, data.get("ground_rules", []))
    if data.get("start_here"):
        render_start_here(doc, data["start_here"])
    render_coverage(doc, data["coverage"])
    if data.get("ratings_legend"):
        render_ratings_legend(doc, data["ratings_legend"])
    if data.get("part_a"):
        render_part_a(doc, data["part_a"])
    if data.get("part_b"):
        render_part_b(doc, data["part_b"])
    if data.get("boundaries"):
        render_boundaries(doc, data["boundaries"])
    render_sign_off(doc, data["sign_off"])
    if data.get("disposition_register"):
        render_disposition_register(doc, data["disposition_register"])
    if data.get("appendix"):
        render_appendix(doc, data["appendix"])
    render_footer(doc)
    return doc


def slugify(name: str) -> str:
    keep = []
    for ch in name:
        if ch.isalnum():
            keep.append(ch)
        elif ch in " &+/-":
            keep.append(" ")
    return "_".join("".join(keep).split())


def generate(data_path: Path, output_path: Path) -> None:
    data = json.loads(data_path.read_text(encoding="utf-8"))
    doc = build_document(data)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"✓ Packet generated: {output_path}")


def main(argv):
    here = Path(__file__).resolve().parent
    if argv and argv[0] == "--all":
        data_dir = here / "data"
        out_dir = here / "outputs"
        for data_file in sorted(data_dir.glob("*.json")):
            data = json.loads(data_file.read_text(encoding="utf-8"))
            if "metadata" not in data or "coverage" not in data:
                continue  # reference file, not a packet
            sub = data["metadata"].get("sub_family", data_file.stem)
            out = out_dir / f"HR_R2_SME_Review_Packet_{slugify(sub)}.docx"
            generate(data_file, out)
        return 0

    if not argv:
        print(__doc__)
        return 1

    data_path = Path(argv[0])
    if len(argv) > 1:
        output_path = Path(argv[1])
    else:
        data = json.loads(data_path.read_text(encoding="utf-8"))
        sub = data["metadata"].get("sub_family", data_path.stem)
        output_path = here / "outputs" / f"HR_R2_SME_Review_Packet_{slugify(sub)}.docx"
    generate(data_path, output_path)
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
