"""Rendering engine for the HR Round 2 SME Review Packet (python-docx).

This module owns all *stable* content (boilerplate that is identical across
specializations) and the mechanics of building a Cargill-branded Word document.
The *variable* content — competencies, essential functions, crosswalk — is
supplied by ``hr_content.py`` and passed in as a ``spec`` dict.

The document structure, section order, and branding follow the
``sme-validation-package`` skill's ``document-structure.md`` and
``stable-content.md`` references, adapted for the Human Resources job family and
the Round 2 (R2) SME feedback round.
"""

from __future__ import annotations

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

# --------------------------------------------------------------------------- #
# Brand constants (Cargill). Hex values without the leading '#'.
# Primary Cargill fonts (Big Caslon / Helvetica Now) are not available in the
# generation environment, so we use the documented fallbacks (Georgia / Arial).
# --------------------------------------------------------------------------- #
LEAF_GREEN = "00843D"
WHITE_GREEN = "F5F9ED"
DARK_GRAY = "333333"
MED_GRAY = "666666"
LIGHT_GRAY = "CCCCCC"
WHITE = "FFFFFF"

HEADING_FONT = "Georgia"
BODY_FONT = "Arial"

ROUND_LABEL = "Round 2 — SME Review"
PREPARED_BY = "Assessment, Competency & Career Framework Team"


def _rgb(hex_str: str) -> RGBColor:
    return RGBColor.from_string(hex_str)


# --------------------------------------------------------------------------- #
# Low-level helpers
# --------------------------------------------------------------------------- #


def _shade_cell(cell, hex_fill: str) -> None:
    """Apply a solid background fill to a table cell."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_run(run, *, font=BODY_FONT, size=11, color=DARK_GRAY, bold=False, italic=False) -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = _rgb(color)
    run.font.bold = bold
    run.font.italic = italic


def _add_run(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    _set_run(run, **kwargs)
    return run


def _green_rule(paragraph, color: str = LEAF_GREEN, size: int = 6) -> None:
    """Draw a horizontal rule as a bottom border on the given paragraph."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _set_table_borders(table, color: str = LIGHT_GRAY, size: int = 4) -> None:
    """Apply uniform single borders to every edge of a table."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)


def _heading(doc, text, *, level=1, page_break=False):
    para = doc.add_paragraph()
    if page_break:
        para.paragraph_format.page_break_before = True
    if level == 1:
        _add_run(para, text, font=HEADING_FONT, size=20, color=LEAF_GREEN, bold=True)
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(12)
        _green_rule(para)
    elif level == 2:
        _add_run(para, text, font=BODY_FONT, size=13, color=DARK_GRAY, bold=True)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
    else:  # level 3
        _add_run(para, text, font=BODY_FONT, size=11, color=DARK_GRAY, bold=True)
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(4)
    return para


def _body(doc, text, *, italic=False, color=DARK_GRAY, size=11):
    para = doc.add_paragraph()
    _add_run(para, text, font=BODY_FONT, size=size, color=color, italic=italic)
    para.paragraph_format.space_after = Pt(6)
    return para


def _bullet(doc, text, *, number=None):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_after = Pt(4)
    prefix = f"{number}. " if number is not None else "— "
    _add_run(para, prefix + text, font=BODY_FONT, size=11, color=DARK_GRAY)
    return para


# --------------------------------------------------------------------------- #
# Page setup, header, footer
# --------------------------------------------------------------------------- #


def _configure_page(doc) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)


def _add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    _set_run(run, font=BODY_FONT, size=8, color=MED_GRAY)


def _build_header_footer(doc, title: str) -> None:
    section = doc.sections[0]
    # Suppress header/footer on the cover (first) page.
    section.different_first_page_header_footer = True

    header = section.header
    h_para = header.paragraphs[0]
    h_para.text = ""
    tab_stops = h_para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_run(h_para, title, font=BODY_FONT, size=8, color=MED_GRAY)
    _add_run(h_para, "\tCargill", font=HEADING_FONT, size=10, color=LEAF_GREEN, bold=True)

    footer = section.footer
    f_para = footer.paragraphs[0]
    f_para.text = ""
    tab_stops = f_para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_run(f_para, "Cargill Confidential — Internal Use Only", font=BODY_FONT, size=8, color=MED_GRAY)
    _add_run(f_para, "\tPage ", font=BODY_FONT, size=8, color=MED_GRAY)
    _add_page_number_field(f_para)


# --------------------------------------------------------------------------- #
# Styled tables
# --------------------------------------------------------------------------- #


def _header_row(table, labels) -> None:
    cells = table.rows[0].cells
    for cell, label in zip(cells, labels):
        _shade_cell(cell, LEAF_GREEN)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        _add_run(para, label, font=BODY_FONT, size=10, color=WHITE, bold=True)


def _data_cell(cell, text, *, size=10, bold=False, color=DARK_GRAY, center=False):
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(para, text, font=BODY_FONT, size=size, color=color, bold=bold)


def _rating_scale_table(doc, rows) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.allow_autofit = False
    _set_table_borders(table)
    _header_row(table, ["Rating", "Label", "Definition"])
    for i, (rating, label, definition) in enumerate(rows):
        row = table.add_row().cells
        if i % 2 == 1:
            for c in row:
                _shade_cell(c, WHITE_GREEN)
        _data_cell(row[0], rating, bold=True, center=True)
        _data_cell(row[1], label, bold=True)
        _data_cell(row[2], definition)
    table.columns[0].width = Inches(0.8)
    table.columns[1].width = Inches(1.7)
    table.columns[2].width = Inches(4.0)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


LEVEL_LABELS = {
    "L4": "L4 — Expert / Thought Leader",
    "L3": "L3 — Advanced Practitioner",
    "L2": "L2 — Practitioner",
    "L1": "L1 — Foundational",
}


def _competency_indicator_table(doc, levels) -> None:
    """Render the L4→L1 behavioral-indicator table for one competency."""
    table = doc.add_table(rows=1, cols=2)
    table.allow_autofit = False
    _set_table_borders(table)
    _header_row(table, ["Proficiency Level", "Behavioral Indicators"])

    shade = False
    for key in ("L4", "L3", "L2", "L1"):
        indicators = levels[key]
        for idx, indicator in enumerate(indicators):
            row = table.add_row().cells
            level_cell, ind_cell = row[0], row[1]
            if idx == 0:
                _shade_cell(level_cell, LEAF_GREEN)
                _data_cell(level_cell, LEVEL_LABELS[key], color=WHITE, bold=True)
            else:
                _shade_cell(level_cell, LEAF_GREEN)
                level_cell.paragraphs[0].add_run("")  # keep band visually contiguous
            if shade:
                _shade_cell(ind_cell, WHITE_GREEN)
            _data_cell(ind_cell, indicator)
            shade = not shade
        shade = False  # reset alternation at each level boundary

    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(5.0)


def _crosswalk_table(doc, competencies, ef_count, crosswalk) -> None:
    headers = ["Competency"] + [f"EF{i}" for i in range(1, ef_count + 1)]
    table = doc.add_table(rows=1, cols=len(headers))
    table.allow_autofit = False
    _set_table_borders(table)
    _header_row(table, headers)
    for i, comp in enumerate(competencies):
        row = table.add_row().cells
        if i % 2 == 1:
            for c in row:
                _shade_cell(c, WHITE_GREEN)
        _data_cell(row[0], comp["name"], bold=True, color=LEAF_GREEN)
        marked = set(crosswalk.get(comp["name"], []))
        for ef_num in range(1, ef_count + 1):
            _data_cell(row[ef_num], "X" if ef_num in marked else "", center=True, bold=True)
    table.columns[0].width = Inches(2.6)
    for col in range(1, len(headers)):
        table.columns[col].width = Inches(0.65)


def _rating_box(doc) -> None:
    table = doc.add_table(rows=2, cols=1)
    _set_table_borders(table, color=LEAF_GREEN, size=6)
    top = table.rows[0].cells[0]
    _shade_cell(top, WHITE_GREEN)
    _data_cell(
        top,
        "SME Rating —   Relevance (1-5): ____      Clarity (1-3): ____      "
        "Level Differentiation (1-3): ____",
        bold=True,
    )
    bottom = table.rows[1].cells[0]
    _data_cell(bottom, "Comments: ________________________________________________________________")
    doc.add_paragraph().paragraph_format.space_after = Pt(12)


# --------------------------------------------------------------------------- #
# Cover page
# --------------------------------------------------------------------------- #


def _build_cover(doc, spec, job_family, session_date) -> None:
    for _ in range(2):
        doc.add_paragraph()

    rule = doc.add_paragraph()
    _green_rule(rule, size=12)
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(title, "Technical Competency Validation Package", font=HEADING_FONT, size=26, color=LEAF_GREEN, bold=True)

    spec_para = doc.add_paragraph()
    spec_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(spec_para, spec["name"], font=BODY_FONT, size=18, color=DARK_GRAY, bold=True)

    fam_para = doc.add_paragraph()
    fam_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(fam_para, job_family, font=BODY_FONT, size=14, color=MED_GRAY)

    round_para = doc.add_paragraph()
    round_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(round_para, ROUND_LABEL, font=BODY_FONT, size=12, color=LEAF_GREEN, bold=True)

    rule2 = doc.add_paragraph()
    rule2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _green_rule(rule2, size=12)

    for _ in range(2):
        doc.add_paragraph()

    for line in (
        "Prepared for: Subject Matter Expert Review Panel",
        f"Prepared by: {PREPARED_BY}",
        f"Date: {session_date}",
        "Classification: Cargill Confidential — Internal Use Only",
    ):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(para, line, font=BODY_FONT, size=11, color=DARK_GRAY)

    for _ in range(2):
        doc.add_paragraph()

    notice = doc.add_table(rows=1, cols=1)
    _set_table_borders(notice, color=LIGHT_GRAY, size=4)
    cell = notice.rows[0].cells[0]
    para = cell.paragraphs[0]
    _add_run(
        para,
        "This document contains proprietary competency architecture materials developed by "
        "Cargill's Assessment, Competency & Career Framework team. Distribution is limited to "
        "designated Subject Matter Expert reviewers and project stakeholders. Do not forward, "
        "copy, or share outside the approved reviewer panel without written authorization from "
        "the project owner.",
        font=BODY_FONT,
        size=9,
        color=MED_GRAY,
        italic=True,
    )


# --------------------------------------------------------------------------- #
# Body sections
# --------------------------------------------------------------------------- #


def _section_overview(doc, spec, job_family) -> None:
    name = spec["name"]
    _heading(doc, "1. Project Overview", level=1, page_break=True)

    _heading(doc, "Purpose Statement", level=2)
    _body(
        doc,
        "The Technical Competency Builder (TCB) initiative is developing technical competencies "
        "across Cargill's 15 job families to support talent decisions spanning selection, "
        "development, performance management, and succession planning. These competencies form "
        "the third layer of Cargill's three-tier competency architecture.",
    )
    _body(
        doc,
        f"Your role as a Subject Matter Expert (SME) is to evaluate whether the draft technical "
        f"competencies accurately represent the critical knowledge, skills, and behaviors required "
        f"for effective performance in the {name} specialization within {job_family}.",
    )
    _body(
        doc,
        "This is the second SME validation round (Round 2). The competencies presented here "
        "incorporate revisions made in response to Round 1 feedback. Where Round 1 raised "
        "questions about wording, level differentiation, or scope, those items have been "
        "revised and are brought back for your confirmation.",
        italic=True,
    )

    _heading(doc, "What We Are Asking You to Do", level=2)
    _body(doc, "During this 60-minute structured session, you will:")
    _bullet(doc, "Review each draft technical competency and its behavioral indicators across four proficiency levels (L1 through L4).", number=1)
    _bullet(doc, f"Rate each competency on relevance to {name}, clarity of indicators, and level differentiation.", number=2)
    _bullet(doc, "Identify gaps: are there critical technical capabilities missing from the draft set?", number=3)
    _bullet(doc, "Validate the mapping between competencies and essential job functions for the Manager II and Advisor bands.", number=4)
    _body(
        doc,
        "Your expert judgment directly shapes the final competency framework. This is not a test — "
        "there are no right or wrong answers. We need your candid professional assessment based on "
        "your experience performing and supervising this work.",
    )


def _section_protocol(doc) -> None:
    _heading(doc, "2. Focus Group Protocol", level=1, page_break=True)
    _heading(doc, "Session Structure", level=2)
    _body(doc, "The 60-minute session follows this structure:")
    for text in (
        "Minutes 0-5: Welcome, introductions, and overview of the session purpose and ground rules.",
        "Minutes 5-15: Walkthrough of the competency architecture and explanation of the rating scales.",
        "Minutes 15-50: Competency-by-competency review, discussion, and individual ratings.",
        "Minutes 50-55: Gap analysis — open discussion on whether any critical technical capabilities are missing.",
        "Minutes 55-60: Wrap-up, next steps, and timeline for incorporating feedback.",
    ):
        _bullet(doc, text)

    _heading(doc, "Ground Rules", level=2)
    for text in (
        "Every perspective is valuable. Disagreement among SMEs is expected and informative.",
        "Focus on the work, not the person. We are evaluating whether behaviors are accurately described.",
        "Specificity helps. Concrete rewording suggestions are more actionable than general reactions.",
        "Silence is data. If a competency generates no discussion, that signals clarity and consensus.",
        "We will capture all feedback, including written feedback via the companion form.",
    ):
        _bullet(doc, text)


def _section_rating_scales(doc, name) -> None:
    _heading(doc, "3. Rating Scales", level=1, page_break=True)

    _heading(doc, "Relevance Scale (5-point)", level=2)
    _body(doc, f"How important is this competency to effective job performance in {name}?")
    _rating_scale_table(
        doc,
        [
            ("5", "Essential", "This competency is critical. Deficiency would result in unacceptable job performance."),
            ("4", "Important", "This competency significantly contributes to effective performance. Most incumbents need it."),
            ("3", "Useful", "This competency is helpful but not critical. Some incumbents may perform adequately without it."),
            ("2", "Minor relevance", f"This competency has limited applicability to {name} work."),
            ("1", "Not relevant", f"This competency does not apply to {name}."),
        ],
    )

    _heading(doc, "Clarity Scale (3-point)", level=2)
    _body(doc, "Are the behavioral indicators clearly written and unambiguous?")
    _rating_scale_table(
        doc,
        [
            ("3", "Clear", "The indicators are specific, observable, and would be interpreted consistently by different raters."),
            ("2", "Partially clear", "The general intent is understandable, but some indicators are vague. Suggest specific rewording."),
            ("1", "Unclear", "The indicators are too abstract or ambiguous to be useful. Significant rewriting needed."),
        ],
    )

    _heading(doc, "Level Differentiation Scale (3-point)", level=2)
    _body(doc, "Do the L1-L4 indicators represent meaningfully different levels of proficiency?")
    _rating_scale_table(
        doc,
        [
            ("3", "Well differentiated", "Each level describes qualitatively different behavior; a trained observer could reliably distinguish levels."),
            ("2", "Partially differentiated", "Some levels are distinct, but others overlap or differ only in degree. Suggest where the gaps are."),
            ("1", "Not differentiated", "The levels read as minor variations of the same behavior. Fundamental restructuring needed."),
        ],
    )


def _section_architecture(doc, name) -> None:
    _heading(doc, "4. Competency Architecture Reference", level=1, page_break=True)

    _heading(doc, "Three-Layer Model Overview", level=2)
    _body(doc, "Cargill's competency architecture operates on three integrated layers:")
    _bullet(doc, "Layer 1 — Values & Behaviors: enterprise-wide expectations that define how we work at Cargill, applied universally.")
    _bullet(doc, "Layer 2 — Common Competencies: cross-functional capabilities (e.g., business acumen, stakeholder management) spanning multiple job families.")
    _bullet(doc, f"Layer 3 — Technical Competencies (under review today): the knowledge, skills, and behaviors that distinguish a strong {name} professional specifically.")

    _heading(doc, "Proficiency Level Definitions", level=2)
    _rating_scale_table(
        doc,
        [
            ("L1", "Foundational", "Applies core technical knowledge under guidance; handles routine matters within established frameworks."),
            ("L2", "Practitioner", "Independently executes the full range of standard technical work; resolves moderately complex issues."),
            ("L3", "Advanced Practitioner", "Handles complex, ambiguous, or high-stakes matters with minimal oversight; mentors others; a go-to expert."),
            ("L4", "Expert / Thought Leader", "Defines and drives technical strategy across the enterprise or externally; navigates unprecedented matters."),
        ],
    )

    _heading(doc, "Band Mapping Reference", level=2)
    _body(doc, "The two primary bands under review today are:")
    _bullet(doc, "Manager II (JL4 equivalent): expected to operate at L3-L4 across most technical competencies.")
    _bullet(doc, "Advisor (senior individual contributor): expected to operate at L3-L4 with deep technical depth.")
    _body(doc, "Essential functions for each band are provided in Section 6 for your reference during the competency review.")


def _section_competencies(doc, spec) -> None:
    competencies = spec["competencies"]
    total = len(competencies)
    _heading(doc, "5. Technical Competencies Under Review", level=1, page_break=True)
    _body(
        doc,
        f"The {total} draft technical competencies for {spec['name']} are presented below. For each "
        "competency, review the definition and the behavioral indicators at all four proficiency "
        "levels, then record your ratings in the box that follows.",
    )

    for i, comp in enumerate(competencies, start=1):
        rule = doc.add_paragraph()
        rule.paragraph_format.page_break_before = i > 1
        rule.paragraph_format.space_before = Pt(6)
        _green_rule(rule)

        head = doc.add_paragraph()
        _add_run(head, f"Competency {i} of {total}: {comp['name']}", font=BODY_FONT, size=14, color=DARK_GRAY, bold=True)
        head.paragraph_format.space_after = Pt(4)

        defn = doc.add_paragraph()
        _add_run(defn, "Definition: ", font=BODY_FONT, size=11, color=DARK_GRAY, bold=True, italic=True)
        _add_run(defn, comp["definition"], font=BODY_FONT, size=11, color=DARK_GRAY, italic=True)
        defn.paragraph_format.space_after = Pt(8)

        _competency_indicator_table(doc, comp["levels"])
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        _rating_box(doc)


def _section_crosswalk(doc, spec, job_family) -> None:
    _heading(doc, "6. Essential Functions Crosswalk", level=1, page_break=True)

    _heading(doc, "Essential Functions by Band", level=2)
    efs = spec["essential_functions"]
    for band in ("Manager II", "Advisor"):
        _heading(doc, f"{band} Essential Functions", level=3)
        for n, ef in enumerate(efs[band], start=1):
            _bullet(doc, ef, number=n)

    _heading(doc, "Competency-to-Essential-Function Mapping", level=2)
    ef_count = len(efs["Manager II"])
    _body(
        doc,
        "The matrix below maps each competency to the Manager II essential functions it supports. "
        'An "X" indicates the competency is exercised in performing that essential function.',
    )
    _crosswalk_table(doc, spec["competencies"], ef_count, spec["crosswalk"])

    _heading(doc, "Essential Function Key (Manager II)", level=3)
    for n, ef in enumerate(efs["Manager II"], start=1):
        _bullet(doc, f"EF{n}: {ef}")

    _heading(doc, "SME Prompt", level=2)
    _body(doc, "Review the mapping above. For each competency:")
    _bullet(doc, "Is the mapping accurate? Are there essential functions that should be linked or unlinked?")
    _bullet(doc, "Are there essential functions not covered by any competency? If so, what technical capability is missing?")


def _section_methodology(doc) -> None:
    _heading(doc, "7. Validation Methodology Note", level=1, page_break=True)
    _body(
        doc,
        "The validation approach follows established content-validity methodology as specified by "
        "the SIOP Principles for the Validation and Use of Personnel Selection Procedures and the "
        "APA Standards for Educational and Psychological Testing.",
    )
    _body(
        doc,
        "Content validity evidence is established when qualified Subject Matter Experts confirm that "
        "the competency content is representative of the critical requirements of the job. Your "
        "individual ratings will be aggregated to compute a Content Validity Ratio (CVR) for each "
        "competency, following Lawshe's (1975) methodology. Competencies meeting the CVR threshold "
        "for your panel size will be retained; those falling below will be revised or removed.",
    )


def _section_confidentiality(doc, name) -> None:
    _heading(doc, "8. Confidentiality and Data Handling", level=1, page_break=True)
    _body(
        doc,
        "All materials in this package are Cargill Confidential. Your individual ratings and comments "
        "will be aggregated with other SME responses. Individual responses will not be attributed by "
        "name in any reporting unless you provide explicit written consent.",
    )
    _body(doc, "The aggregated validation data will be used to:")
    _bullet(doc, f"Finalize the technical competency framework for {name}.")
    _bullet(doc, "Compute content-validity statistics for psychometric documentation.")
    _bullet(doc, "Inform revisions to behavioral indicators based on SME feedback.")
    _bullet(doc, "Support the overall TCB initiative's audit and governance requirements.")
    _body(
        doc,
        "Please return all physical copies of this document at the end of the session or destroy them "
        "per Cargill's information handling policy.",
    )


def _section_contact(doc) -> None:
    _heading(doc, "9. Contact Information", level=1, page_break=True)
    for line in (
        "Project Owner:",
        "Christopher Honts, Ph.D.",
        "Manager, Assessment, Competency & Career Framework",
        "Cargill — Human Resources",
        "Wayzata, Minnesota",
    ):
        _body(doc, line)
    _body(
        doc,
        "For questions about this validation package, the focus group session, or the TCB initiative, "
        "please contact the project owner directly.",
    )


def _appendix_specialization(doc, spec, job_family) -> None:
    _heading(doc, f"Appendix A: Specialization Overview — {spec['name']}", level=1, page_break=True)
    _body(doc, spec["description"])
    _body(
        doc,
        f"This overview is provided so reviewers share a common understanding of the {spec['name']} "
        f"specialization within {job_family} before evaluating the draft competencies.",
        italic=True,
    )


def _appendix_glossary(doc) -> None:
    _heading(doc, "Appendix B: Glossary of Terms", level=1, page_break=True)
    terms = [
        ("Behavioral Indicator", "A specific, observable action that demonstrates proficiency in a competency at a given level."),
        ("Competency", "A cluster of related knowledge, skills, and behaviors that differentiate levels of job performance."),
        ("Content Validity", "Evidence that a framework adequately represents the critical requirements of the job, as judged by SMEs."),
        ("Content Validity Ratio (CVR)", "A statistical index (Lawshe, 1975) quantifying SME agreement that a competency is essential."),
        ("Essential Function", "A fundamental duty of a job, as defined in the job description, distinguished from marginal functions."),
        ("Job Level (JL)", "Cargill's career band designation (JL1 through JL4) reflecting scope, complexity, and impact."),
        ("Observable-Measurable-Discriminant (OMD)", "Quality criteria for indicators: observable, measurable, and discriminant across levels."),
        ("Proficiency Level (L1-L4)", "A qualitative tier of capability within a technical competency, from Foundational to Expert."),
        ("Subject Matter Expert (SME)", "An individual with deep experience in the specialization, qualified to evaluate competency content."),
        ("Technical Competency", "A Layer 3 competency capturing specialized knowledge and behaviors unique to a job family or specialization."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.allow_autofit = False
    _set_table_borders(table)
    _header_row(table, ["Term", "Definition"])
    for i, (term, definition) in enumerate(terms):
        row = table.add_row().cells
        if i % 2 == 1:
            for c in row:
                _shade_cell(c, WHITE_GREEN)
        _data_cell(row[0], term, bold=True, color=LEAF_GREEN)
        _data_cell(row[1], definition)
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.4)


# --------------------------------------------------------------------------- #
# Public entry point
# --------------------------------------------------------------------------- #


def build_packet(spec: dict, job_family: str, session_date: str, out_path: str) -> str:
    """Build one HR Round 2 SME Review Packet and save it to ``out_path``.

    Returns the path written.
    """
    doc = Document()
    _configure_page(doc)
    title = f"HR R2 SME Review Packet — {spec['name']}"
    _build_header_footer(doc, title)

    _build_cover(doc, spec, job_family, session_date)
    _section_overview(doc, spec, job_family)
    _section_protocol(doc)
    _section_rating_scales(doc, spec["name"])
    _section_architecture(doc, spec["name"])
    _section_competencies(doc, spec)
    _section_crosswalk(doc, spec, job_family)
    _section_methodology(doc)
    _section_confidentiality(doc, spec["name"])
    _section_contact(doc)
    _appendix_specialization(doc, spec, job_family)
    _appendix_glossary(doc)

    doc.save(out_path)
    return out_path
