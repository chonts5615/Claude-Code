"""
Word document (.docx) content extractor.
"""

import io
from pathlib import Path
from uuid import uuid4

from src.extractors.base import BaseExtractor
from src.schemas.content import (
    ContentBlock,
    ContentType,
    DocumentMetadata,
    ExtractedContent,
)


class DocxExtractor(BaseExtractor):
    """Extract content from Microsoft Word documents."""

    def supported_extensions(self) -> list[str]:
        return [".docx", ".doc"]

    def extract(self, file_path: Path) -> ExtractedContent:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        doc = Document(str(file_path))
        blocks: list[ContentBlock] = []
        warnings: list[str] = []
        has_images = False
        has_tables = False
        word_count = 0

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            word_count += len(text.split())
            block_id = f"block_{uuid4().hex[:8]}"

            # Determine content type from style
            style_name = para.style.name if para.style else ""

            if "Heading" in style_name or "Title" in style_name:
                level = self._parse_heading_level(style_name)
                blocks.append(ContentBlock(
                    block_id=block_id,
                    content_type=ContentType.HEADING,
                    text=text,
                    level=level,
                ))
            elif "List Bullet" in style_name or "List" in style_name:
                # Check if we can merge with previous bullet list
                if blocks and blocks[-1].content_type == ContentType.BULLET_LIST:
                    blocks[-1].items.append(text)
                else:
                    blocks.append(ContentBlock(
                        block_id=block_id,
                        content_type=ContentType.BULLET_LIST,
                        items=[text],
                    ))
            elif "List Number" in style_name:
                if blocks and blocks[-1].content_type == ContentType.NUMBERED_LIST:
                    blocks[-1].items.append(text)
                else:
                    blocks.append(ContentBlock(
                        block_id=block_id,
                        content_type=ContentType.NUMBERED_LIST,
                        items=[text],
                    ))
            else:
                blocks.append(ContentBlock(
                    block_id=block_id,
                    content_type=ContentType.PARAGRAPH,
                    text=text,
                ))

            # Check for inline images
            for run in para.runs:
                if run._element.findall(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                ):
                    has_images = True

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            has_tables = True
            rows = []
            headers = []

            for row_idx, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                if row_idx == 0:
                    headers = row_data
                else:
                    rows.append(row_data)
                    for cell_text in row_data:
                        word_count += len(cell_text.split())

            block_id = f"table_{uuid4().hex[:8]}"
            blocks.append(ContentBlock(
                block_id=block_id,
                content_type=ContentType.TABLE,
                table_data=rows,
                table_headers=headers,
                metadata={"table_index": table_idx},
            ))

        # Extract images from document relationships
        image_blocks = self._extract_images(doc, warnings)
        if image_blocks:
            has_images = True
            blocks.extend(image_blocks)

        # Build metadata
        core_props = doc.core_properties
        metadata = DocumentMetadata(
            title=core_props.title or self._infer_title(blocks),
            author=core_props.author,
            word_count=word_count,
            has_images=has_images,
            has_tables=has_tables,
        )

        return ExtractedContent(
            source_file=str(file_path),
            source_format="docx",
            blocks=blocks,
            metadata=metadata,
            extraction_warnings=warnings,
        )

    def _parse_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        if "Title" in style_name:
            return 1
        for i in range(1, 7):
            if str(i) in style_name:
                return i
        return 1

    def _infer_title(self, blocks: list[ContentBlock]) -> str:
        """Infer document title from first heading."""
        for block in blocks:
            if block.content_type == ContentType.HEADING and block.level <= 2:
                return block.text or ""
        return ""

    def _extract_images(self, doc, warnings: list[str]) -> list[ContentBlock]:
        """Extract embedded images from the document."""
        image_blocks = []
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        image_data = rel.target_part.blob
                        content_type = rel.target_part.content_type
                        ext = content_type.split("/")[-1] if content_type else "png"

                        block_id = f"img_{uuid4().hex[:8]}"
                        image_blocks.append(ContentBlock(
                            block_id=block_id,
                            content_type=ContentType.IMAGE,
                            image_data=image_data,
                            image_format=ext,
                        ))
                    except Exception as e:
                        warnings.append(f"Failed to extract image: {e}")
        except Exception as e:
            warnings.append(f"Image extraction error: {e}")

        return image_blocks
